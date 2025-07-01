from docx import Document
import re
import unicodedata

def extrair_texto_txt(caminho_txt):
    with open(caminho_txt, "r", encoding="utf-8") as arquivo:
        return arquivo.read()

def extrair_bloco(texto, titulo):
    padrao = rf'{titulo}(.*?)(\n[A-Z][^\n]+\n|\Z)'
    resultado = re.search(padrao, texto, re.DOTALL | re.IGNORECASE)
    return resultado.group(1).strip() if resultado else ''

def buscar_numero(texto, padrao):
    resultado = re.search(padrao, texto, re.IGNORECASE)
    if resultado:
        valor = resultado.group(1).replace('.', '').replace(',', '.')
        try:
            return float(valor)
        except:
            return 0.0
    return 0.0

def buscar_patrimonios(texto):
    """Busca todos os valores de bens e direitos por ano"""
    matches = re.findall(r'Bens e direitos em 31/12/(\d{4})\s*([\d\.,]+)', texto)
    patrimonios = {}
    for ano, valor in matches:
        valor_float = float(valor.replace('.', '').replace(',', '.'))
        patrimonios[ano] = valor_float
    return patrimonios

def gerar_relatorio(texto, nome_arquivo=None):
    doc = Document()

    # Nome
    nome_match = re.search(r'Nome:\s*([^\n\r]+)', texto, re.IGNORECASE)
    nome = nome_match.group(1).strip().title() if nome_match else "Nao_Encontrado"

    # CPF
    cpf_match = re.search(r'CPF:\s*([\d\.\-]+)', texto)
    cpf = cpf_match.group(1) if cpf_match else "Não encontrado"

    # Data de nascimento
    nasc_match = re.search(r'Data de Nascimento:\s*([^\n\r]+)', texto)
    nascimento = nasc_match.group(1).strip() if nasc_match else "Não encontrada"

    # Dependentes
    dependentes = buscar_numero(texto, r'Total de Dependentes.*?(\d+)')

    # Nome do arquivo baseado no nome do cliente
    if not nome_arquivo:
        nome_sem_acentos = unicodedata.normalize('NFKD', nome).encode('ASCII', 'ignore').decode('utf-8')
        nome_limpo = re.sub(r'[^\w\-]', '_', nome_sem_acentos).strip('_')
        nome_arquivo = f"Relatorio_{nome_limpo}.docx"

    # Patrimônio por ano
    patrimonios = buscar_patrimonios(texto)
    anos_disponiveis = sorted(patrimonios.keys(), reverse=True)
    patrimonio_ano_atual = patrimonios[anos_disponiveis[0]] if anos_disponiveis else 0.0
    patrimonio_ano_anterior = patrimonios[anos_disponiveis[1]] if len(anos_disponiveis) > 1 else None
    evolucao_patrimonial = patrimonio_ano_atual - patrimonio_ano_anterior if patrimonio_ano_anterior else None

    # Dívidas
    dividas_valor = buscar_numero(texto, r'DÍVIDAS E ÔNUS REAIS\s*\n.*?Total.*?([\d\.,]+)')
    possui_dividas = "Não possui" if "DÍVIDAS E ÔNUS REAIS\n\nSem Informações" in texto else f"Possui - R$ {dividas_valor:,.2f}"

    # Atividade Rural
    bloco_rural = extrair_bloco(texto, 'DEMONSTRATIVO DE ATIVIDADE RURAL')
    receita_rural = buscar_numero(bloco_rural, r'Receita bruta total\s*([\d\.,]+)')
    if receita_rural == 0:
        receita_rural = buscar_numero(bloco_rural, r'Receita Bruta da Atividade Rural.*?([\d\.,]+)')
    resultado_rural = buscar_numero(bloco_rural, r'Resultado Tributável\s*([\d\.,]+)')
    resultado_nao_tributavel = buscar_numero(bloco_rural, r'Resultado Não Tributável\s*([\d\.,]+)')

    # Se não achou nos blocos, busca no texto geral
    if receita_rural == 0:
        receita_rural = buscar_numero(texto, r'Receita bruta total\s*([\d\.,]+)')
    if resultado_rural == 0:
        resultado_rural = buscar_numero(texto, r'Resultado Tributável\s*([\d\.,]+)')

    # Rendimentos gerais
    rendimentos_tributaveis = buscar_numero(texto, r'Rendimentos Tributáveis.*?([\d\.,]+)')
    rendimentos_isentos = buscar_numero(texto, r'Rendimentos Isentos.*?([\d\.,]+)')

    # Limites sugeridos
    limite_sem_garantia = round(resultado_rural * 3, 2)
    limite_com_garantia = round(patrimonio_ano_atual * 0.5, 2)

    # 📝 Monta o relatório
    doc.add_heading(f'Relatório de Análise - {nome}', level=1)

    doc.add_heading('1. Dados do Produtor:', level=2)
    doc.add_paragraph(f'Nome: {nome}')
    doc.add_paragraph(f'CPF: {cpf}')
    doc.add_paragraph(f'Data de Nascimento: {nascimento}')
    doc.add_paragraph(f'Dependentes: {int(dependentes)}')

    doc.add_heading('2. Dados Financeiros:', level=2)
    if patrimonio_ano_atual:
        doc.add_paragraph(f'Patrimônio declarado ({anos_disponiveis[0]}): R$ {patrimonio_ano_atual:,.2f}')
    if patrimonio_ano_anterior:
        doc.add_paragraph(f'Patrimônio anterior ({anos_disponiveis[1]}): R$ {patrimonio_ano_anterior:,.2f}')
        doc.add_paragraph(f'Evolução patrimonial: R$ {evolucao_patrimonial:,.2f}')
    doc.add_paragraph(f'Dívidas Declaradas: {possui_dividas}')
    doc.add_paragraph(f'Receita Bruta da Atividade Rural: R$ {receita_rural:,.2f}')
    doc.add_paragraph(f'Resultado Tributável da Atividade Rural: R$ {resultado_rural:,.2f}')
    doc.add_paragraph(f'Resultado Não Tributável: R$ {resultado_nao_tributavel:,.2f}')
    doc.add_paragraph(f'Rendimentos Tributáveis: R$ {rendimentos_tributaveis:,.2f}')
    doc.add_paragraph(f'Rendimentos Isentos: R$ {rendimentos_isentos:,.2f}')

    doc.add_heading('3. Análise de Capacidade:', level=2)
    doc.add_paragraph(f'Limite sugerido sem garantia: R$ {limite_sem_garantia:,.2f}')
    doc.add_paragraph(f'Limite sugerido com garantia (50% do patrimônio): R$ {limite_com_garantia:,.2f}')

    doc.add_heading('4. Pontos de Atenção:', level=2)
    if patrimonio_ano_atual < 1000000:
        doc.add_paragraph('- Patrimônio considerado modesto para atividade rural.')
    if receita_rural < 500000:
        doc.add_paragraph('- Receita da atividade rural baixa.')
    if resultado_rural < 100000:
        doc.add_paragraph('- Resultado líquido da atividade rural baixo, margens apertadas.')
    if "Não possui" in possui_dividas:
        doc.add_paragraph('- Sem histórico de dívidas recentes, pode impactar avaliação de crédito.')

    doc.add_heading('5. Considerações Finais:', level=2)
    doc.add_paragraph(
        'Este relatório foi gerado automaticamente com base nos dados extraídos da Declaração de Imposto de Renda '
        'do contribuinte. Recomenda-se análise complementar documental e, se necessário, vistoria dos ativos declarados.'
    )

    doc.save(nome_arquivo)
    print(f'Relatório gerado com sucesso: {nome_arquivo}')


# ==============================
# 🚀 USO DO SCRIPT COM TXT
# ==============================

arquivo_txt = "relatorio.txt"  # Altere para o caminho do seu .txt
texto_extraido = extrair_texto_txt(arquivo_txt)
gerar_relatorio(texto_extraido)
